import streamlit as st
import pandas as pd
import altair as alt
from ScriptXLM_RoBERTa import predict_with_loaded_model, predict_with_top_5_laws, predict_with_top_5_words_and_sentences
from io import BytesIO
from docx import Document

# Function to hide Streamlit menu and footer
def hide_streamlit_menu_footer():
    st.markdown(
        """
        <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        .stButton>button {
            width: 100%;
            margin: 0;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

# Function to style buttons
def style_buttons():
    st.markdown(
        """
        <style>
        .red-button {
            background-color: #ff4b4b !important;
            color: white !important;
            padding: 15px;
            border-radius: 5px;
        }
        .green-button {
            background-color: #28a745 !important;
            color: white !important;
            padding: 15px;
            border-radius: 5px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

# Landing page function
def landing_page():
    col1, col2 = st.columns([3, 1])
    with col1:
        st.title("Objection Helper")
        st.markdown("Welcome to the Objection Helper. Click start to proceed.")
    with col2:
        st.image("images/GASD_1.png", use_container_width=True)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        if st.button("Start", key="start_button", use_container_width=True):
            st.session_state.page = "Submit Information"

# Submit information page function
def submit_information_page():
    st.title("Submit Information")

    # Initialize session state variables if not set
    if "id_input" not in st.session_state:
        st.session_state.id_input = ""
    if "subject_input" not in st.session_state:
        st.session_state.subject_input = "Waste Fines"
    if "objection_input" not in st.session_state:
        st.session_state.objection_input = ""
    if "proceed_clicked" not in st.session_state:
        st.session_state.proceed_clicked = False

    with st.container():
        col1, col2 = st.columns(2)
        with col1:
            st.text_input("ID of Objection*", value=st.session_state.id_input, key="id_input")
            st.selectbox(
                "Subject*",
                ["Waste Fines", "Vehicle Towing"],
                index=["Waste Fines", "Vehicle Towing"].index(st.session_state.subject_input),
                key="subject_input"
            )
            st.text_area("Objection*", value=st.session_state.objection_input, key="objection_input")
            
        with col2:
            st.image("images/GASD_1.png", use_container_width=True)
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Quit", key="quit_button", use_container_width=True):
                st.session_state.page = "Landing"
        with col2:
            if st.button("Proceed", key="proceed_button", use_container_width=True):
                st.session_state.proceed_clicked = True

    if st.session_state.proceed_clicked:
        if st.session_state.objection_input:
            prediction = predict_with_loaded_model(st.session_state.objection_input)
            st.session_state.result = prediction
            st.session_state.page = "Result"
            st.session_state.proceed_clicked = False  # Reset to avoid re-execution
        else:
            st.error("Please enter the objection details before proceeding.")

# Result page function
def result_page():
    st.title("Result")

    result = st.session_state.get("result", ("No result available", [0, 0]))
    predicted_label, laws, probabilities = result
    label_map = {0: "Ongegrond", 1: "Gegrond"}

    label_text = label_map.get(predicted_label, "Unknown")
    st.subheader(f"Prediction: {label_text}")
    st.write(f"Probability: {probabilities[predicted_label]:.2f}")

    # Store objection input in session state
    st.text_input("Objection ID", value=st.session_state.id_input, key="id_input", disabled=True)
    st.text_input("Subject", value=st.session_state.subject_input, key="subject_input", disabled=True)
    st.write("Objection")
    st.markdown(f"<p>{st.session_state.objection_input}</p>", unsafe_allow_html=True)

    # Create and download result document
    def create_result_doc():
        doc = Document()
        doc.add_heading("Objection Details", 0)
        doc.add_paragraph(f"Objection ID: {st.session_state.get('id_input', '')}")
        doc.add_paragraph(f"Subject: {st.session_state.get('subject_input', '')}")
        doc.add_paragraph(f"Objection: {st.session_state.get('objection_input', '')}")
        doc.add_heading("Prediction", level=1)
        doc.add_paragraph(f"Prediction: {label_text}")
        doc.add_paragraph(f"Probability: {probabilities[predicted_label]:.2f}")
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Restart", key="restart_button", use_container_width=True):
            st.session_state.clear()
            st.session_state.page = "Landing"
    with col2:
        st.download_button(
            label="Download Result",
            use_container_width=True,
            key="download_button_result",
            data=create_result_doc(),
            file_name="result.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# Initialize session state
if "page" not in st.session_state:
    st.session_state.page = "Landing"

# Apply CSS
hide_streamlit_menu_footer()
style_buttons()

# Page routing
if st.session_state.page == "Landing":
    landing_page()
elif st.session_state.page == "Submit Information":
    submit_information_page()
elif st.session_state.page == "Result":
    result_page()
